"""Text extraction from uploaded resume files (PDF, DOCX).

This module only handles *raw text extraction*. Structured field extraction
(skills, experience, education, ...) is a separate service (Phase 6).
"""
import io

import docx
import fitz  # PyMuPDF

ALLOWED_CONTENT_TYPES = {
    "application/pdf": "pdf",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document": "docx",
}
ALLOWED_EXTENSIONS = {"pdf", "docx"}


class ResumeParsingError(Exception):
    """Raised when a resume file can't be validated or its text extracted."""


def get_file_extension(filename: str) -> str:
    if "." not in filename:
        return ""
    return filename.rsplit(".", 1)[-1].lower()


def validate_resume_file(filename: str, content_type: str | None, size_bytes: int, max_size_mb: int) -> str:
    """Validate an uploaded resume file. Returns the normalized extension ("pdf"/"docx").

    Raises ResumeParsingError with a user-facing message on any validation failure.
    """
    if size_bytes == 0:
        raise ResumeParsingError("The uploaded file is empty.")

    max_size_bytes = max_size_mb * 1024 * 1024
    if size_bytes > max_size_bytes:
        raise ResumeParsingError(f"File is too large. Maximum size is {max_size_mb}MB.")

    extension = get_file_extension(filename)
    if extension not in ALLOWED_EXTENSIONS:
        raise ResumeParsingError("Unsupported file type. Please upload a PDF or DOCX file.")

    return extension


def extract_text_from_pdf(content: bytes) -> str:
    try:
        with fitz.open(stream=content, filetype="pdf") as document:
            if document.is_encrypted:
                raise ResumeParsingError(
                    "This PDF is password-protected. Please upload an unprotected file."
                )
            pages_text = [page.get_text() for page in document]
    except ResumeParsingError:
        raise
    except Exception as exc:
        raise ResumeParsingError(
            "Couldn't read this PDF. It may be corrupted or not a valid PDF file."
        ) from exc

    return "\n".join(pages_text)


def extract_text_from_docx(content: bytes) -> str:
    try:
        document = docx.Document(io.BytesIO(content))
    except Exception as exc:
        raise ResumeParsingError(
            "Couldn't read this DOCX file. It may be corrupted or not a valid Word document."
        ) from exc

    paragraphs = [paragraph.text for paragraph in document.paragraphs]

    # Table cells are common in resumes (e.g. skills laid out in a grid) and
    # python-docx's paragraph iteration skips them entirely.
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    paragraphs.append(cell.text)

    return "\n".join(paragraphs)


def extract_resume_text(filename: str, content_type: str | None, content: bytes, max_size_mb: int) -> str:
    """Validate and extract raw text from an uploaded resume file.

    Raises ResumeParsingError on any failure — callers should turn this into a 400.
    """
    extension = validate_resume_file(filename, content_type, len(content), max_size_mb)

    if extension == "pdf":
        text = extract_text_from_pdf(content)
    else:
        text = extract_text_from_docx(content)

    text = text.strip()
    if not text:
        raise ResumeParsingError(
            "No readable text found in this file. It may be a scanned image "
            "without selectable text — try a different file."
        )

    return text
