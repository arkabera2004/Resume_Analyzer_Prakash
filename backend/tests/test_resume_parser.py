"""Tests for the resume text-extraction service."""
import io

import docx
import fitz
import pytest

from app.services.resume_parser import (
    ResumeParsingError,
    extract_resume_text,
    extract_text_from_docx,
    extract_text_from_pdf,
    get_file_extension,
    validate_resume_file,
)


def make_pdf_bytes(text: str = "Jane Doe\nSoftware Engineer\nPython, React, MongoDB") -> bytes:
    document = fitz.open()
    page = document.new_page()
    page.insert_text((72, 72), text)
    data = document.tobytes()
    document.close()
    return data


def make_docx_bytes(paragraphs: list[str] | None = None) -> bytes:
    paragraphs = paragraphs or ["Jane Doe", "Software Engineer", "Python, React, MongoDB"]
    document = docx.Document()
    for paragraph in paragraphs:
        document.add_paragraph(paragraph)
    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def test_get_file_extension():
    assert get_file_extension("resume.pdf") == "pdf"
    assert get_file_extension("resume.DOCX") == "docx"
    assert get_file_extension("noextension") == ""


def test_validate_resume_file_accepts_pdf_and_docx():
    assert validate_resume_file("resume.pdf", "application/pdf", 1000, max_size_mb=5) == "pdf"
    assert validate_resume_file("resume.docx", "application/msword", 1000, max_size_mb=5) == "docx"


def test_validate_resume_file_rejects_unsupported_extension():
    with pytest.raises(ResumeParsingError, match="Unsupported file type"):
        validate_resume_file("resume.txt", "text/plain", 1000, max_size_mb=5)


def test_validate_resume_file_rejects_empty_file():
    with pytest.raises(ResumeParsingError, match="empty"):
        validate_resume_file("resume.pdf", "application/pdf", 0, max_size_mb=5)


def test_validate_resume_file_rejects_oversized_file():
    ten_mb = 10 * 1024 * 1024
    with pytest.raises(ResumeParsingError, match="too large"):
        validate_resume_file("resume.pdf", "application/pdf", ten_mb, max_size_mb=5)


def test_extract_text_from_pdf_returns_readable_text():
    pdf_bytes = make_pdf_bytes("Jane Doe\nSoftware Engineer")
    text = extract_text_from_pdf(pdf_bytes)
    assert "Jane Doe" in text
    assert "Software Engineer" in text


def test_extract_text_from_pdf_rejects_malformed_file():
    with pytest.raises(ResumeParsingError, match="Couldn't read this PDF"):
        extract_text_from_pdf(b"this is not a real pdf")


def test_extract_text_from_docx_returns_paragraphs_and_table_cells():
    document = docx.Document()
    document.add_paragraph("Jane Doe")
    table = document.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "Python"
    table.rows[0].cells[1].text = "MongoDB"
    buffer = io.BytesIO()
    document.save(buffer)

    text = extract_text_from_docx(buffer.getvalue())
    assert "Jane Doe" in text
    assert "Python" in text
    assert "MongoDB" in text


def test_extract_text_from_docx_rejects_malformed_file():
    with pytest.raises(ResumeParsingError, match="Couldn't read this DOCX"):
        extract_text_from_docx(b"this is not a real docx")


def test_extract_resume_text_end_to_end_pdf():
    pdf_bytes = make_pdf_bytes()
    text = extract_resume_text("resume.pdf", "application/pdf", pdf_bytes, max_size_mb=5)
    assert "Jane Doe" in text


def test_extract_resume_text_end_to_end_docx():
    docx_bytes = make_docx_bytes()
    text = extract_resume_text(
        "resume.docx",
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        docx_bytes,
        max_size_mb=5,
    )
    assert "Jane Doe" in text


def test_extract_resume_text_rejects_pdf_with_no_extractable_text():
    # A blank PDF page has no text at all.
    document = fitz.open()
    document.new_page()
    blank_pdf = document.tobytes()
    document.close()

    with pytest.raises(ResumeParsingError, match="No readable text"):
        extract_resume_text("blank.pdf", "application/pdf", blank_pdf, max_size_mb=5)
